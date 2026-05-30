import os
import sys


# ──────────────────────────────────────────────────────────────────────────────
#  WORD → PDF
#  Strategy: mammoth (DOCX→HTML) → pymupdf story renderer (headless, no Qt)
#  Fallback 1: fpdf2 plain-text extraction
# ──────────────────────────────────────────────────────────────────────────────
def word_to_pdf(docx_path: str, pdf_path: str = None) -> str:
    """Convert a .docx file to PDF.  Returns the path of the created PDF."""
    if pdf_path is None:
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    # ── Primary: mammoth → HTML → pymupdf (fitz) Story ──────────────────────
    try:
        import mammoth
        import fitz  # pymupdf

        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value

        styled_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body  {{ font-family: Arial, sans-serif; font-size: 11pt;
           line-height: 1.5; margin: 0; color: #222; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
  th, td{{ border: 1px solid #555; padding: 6px; text-align: left; }}
  th    {{ background: #eee; font-weight: bold; }}
  img   {{ max-width: 100%; height: auto; }}
  h1,h2,h3 {{ color: #1a2a3a; }}
  p     {{ margin-bottom: 8px; }}
</style>
</head>
<body>{html}</body>
</html>"""

        # pymupdf Story API (available since pymupdf 1.23)
        story = fitz.Story(html=styled_html)
        mediabox = fitz.paper_rect("a4")          # 595 × 842 pts
        margin   = 50
        where    = mediabox + (margin, margin, -margin, -margin)

        writer = fitz.DocumentWriter(pdf_path)
        more   = True
        while more:
            device, more = writer.begin_page(mediabox)
            more, _ = story.place(where)
            story.draw(device)
            writer.end_page()
        writer.close()
        return pdf_path

    except AttributeError:
        # pymupdf < 1.23 — Story API not available, fall through
        pass
    except Exception as e:
        print(f"[word_to_pdf] primary path failed: {e}")

    # ── Fallback 1: mammoth → HTML → Qt QPdfWriter ───────────────────────────
    try:
        import mammoth
        from PySide6.QtGui import QTextDocument, QPdfWriter, QPageLayout, QPageSize
        from PySide6.QtCore import QMarginsF, QSizeF
        from PySide6.QtWidgets import QApplication
        import sys

        # Ensure QApplication exists (needed for Qt painting)
        if not QApplication.instance():
            _app = QApplication(sys.argv)

        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value

        styled_html = f"""<html><head><style>
          body{{font-family:sans-serif;line-height:1.5;color:#333;}}
          table{{border-collapse:collapse;width:100%;margin:10px 0;border:1px solid #444;}}
          th,td{{border:1px solid #444;padding:7px;text-align:left;}}
          img{{max-width:100%;height:auto;}}
          h1,h2,h3{{color:#1a2a3a;}} p{{margin-bottom:9px;}}
        </style></head><body>{html}</body></html>"""

        doc = QTextDocument()
        doc.setHtml(styled_html)
        doc.setPageSize(QSizeF(595, 842))          # A4 in points

        writer = QPdfWriter(pdf_path)
        writer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
        writer.setPageMargins(QMarginsF(15, 15, 15, 15), QPageLayout.Unit.Millimeter)
        doc.print_(writer)
        return pdf_path

    except Exception as e:
        print(f"[word_to_pdf] Qt path failed: {e}")

    # ── Fallback 2: python-docx plain text → fpdf2 ───────────────────────────
    try:
        from docx import Document
        from fpdf import FPDF

        doc = Document(docx_path)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", size=12)

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                try:
                    pdf.multi_cell(0, 8, text)
                except Exception:
                    pdf.multi_cell(0, 8, text.encode("latin-1", "replace").decode("latin-1"))
            else:
                pdf.ln(4)

        pdf.output(pdf_path)
        return pdf_path

    except Exception as e:
        raise RuntimeError(f"word_to_pdf: all conversion paths failed — {e}")


# ──────────────────────────────────────────────────────────────────────────────
#  PDF → WORD
#  Strategy 1: pdf2docx  (layout-aware)
#  Strategy 2: pymupdf text blocks → python-docx paragraphs
#  Strategy 3: pypdf text extraction → python-docx (last resort)
# ──────────────────────────────────────────────────────────────────────────────
def pdf_to_word(pdf_path: str, docx_path: str = None) -> str:
    """Convert a PDF to .docx.  Returns the path of the created DOCX."""
    if docx_path is None:
        docx_path = os.path.splitext(pdf_path)[0] + ".docx"

    # ── Primary: pdf2docx ────────────────────────────────────────────────────
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        # Verify the output isn't empty
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 1000:
            return docx_path
        print("[pdf_to_word] pdf2docx produced an empty/tiny file, trying fallback…")
    except Exception as e:
        print(f"[pdf_to_word] pdf2docx failed: {e}")

    # ── Fallback 1: pymupdf block extraction → python-docx ──────────────────
    try:
        import fitz
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc  = Document()
        pdf  = fitz.open(pdf_path)

        for page_num in range(len(pdf)):
            page   = pdf[page_num]
            blocks = page.get_text("blocks")   # list of (x0,y0,x1,y1,text,…)

            if page_num > 0:
                doc.add_page_break()

            # Sort top-to-bottom, then left-to-right
            blocks.sort(key=lambda b: (round(b[1] / 10), b[0]))

            for block in blocks:
                text = block[4].strip()
                if not text:
                    continue
                para = doc.add_paragraph(text)
                para.paragraph_format.space_after = Pt(4)

        pdf.close()
        doc.save(docx_path)
        return docx_path

    except Exception as e:
        print(f"[pdf_to_word] pymupdf fallback failed: {e}")

    # ── Fallback 2: pypdf plain text ─────────────────────────────────────────
    try:
        from docx import Document
        from pypdf import PdfReader

        doc    = Document()
        reader = PdfReader(pdf_path)

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if i > 0:
                doc.add_page_break()
            for line in text.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())

        doc.save(docx_path)
        return docx_path

    except Exception as e:
        raise RuntimeError(f"pdf_to_word: all conversion paths failed — {e}")


# ──────────────────────────────────────────────────────────────────────────────
#  Other doc utilities (unchanged logic, just kept consistent imports)
# ──────────────────────────────────────────────────────────────────────────────
def image_to_pdf(image_paths: list, pdf_path: str = "output.pdf") -> str:
    import img2pdf
    with open(pdf_path, "wb") as f:
        f.write(img2pdf.convert(image_paths))
    return pdf_path


def split_pdf(pdf_path: str, output_dir: str = None) -> list:
    import pikepdf

    if output_dir is None:
        output_dir = os.path.dirname(pdf_path) or "."
    os.makedirs(output_dir, exist_ok=True)

    base  = os.path.splitext(os.path.basename(pdf_path))[0]
    paths = []

    with pikepdf.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            out     = pikepdf.Pdf.new()
            out.pages.append(page)
            out_path = os.path.join(output_dir, f"{base}_page_{i+1}.pdf")
            out.save(out_path)
            out.close()
            paths.append(out_path)

    return paths


def merge_pdfs(pdf_list: list, output_path: str = "merged.pdf") -> str:
    """Merge a list of PDF paths into one file."""
    from pypdf import PdfWriter
    writer = PdfWriter()
    for p in pdf_list:
        writer.append(p)
    writer.write(output_path)
    writer.close()
    return output_path


def text_to_pdf(text: str, pdf_path: str = "output.pdf") -> str:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=12)

    for line in text.splitlines():
        try:
            pdf.multi_cell(0, 8, line)
        except Exception:
            pdf.multi_cell(0, 8, line.encode("latin-1", "replace").decode("latin-1"))
    pdf.output(pdf_path)
    return pdf_path
